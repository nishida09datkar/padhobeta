from docx import Document
from models.schemas import ParsedDocument


def parse_docx(file_path: str) -> ParsedDocument:
    doc = Document(file_path)

    title = None
    if doc.core_properties.title:
        title = doc.core_properties.title

    lines = []
    page_count = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style_name = para.style.name if para.style else ""

        if "Heading 1" in style_name:
            lines.append(f"# {text}")
        elif "Heading 2" in style_name:
            lines.append(f"## {text}")
        elif "Heading 3" in style_name:
            lines.append(f"### {text}")
        elif "Heading 4" in style_name:
            lines.append(f"#### {text}")
        elif "List Bullet" in style_name:
            lines.append(f"- {text}")
        elif "List Number" in style_name:
            lines.append(f"1. {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        lines.append("")
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    content = "\n".join(lines)

    if not content.strip():
        content = "[DOCX file appears to be empty or contains no readable text.]"

    return ParsedDocument(
        content=content,
        title=title,
        doc_type="docx",
        page_count=page_count,
        metadata={
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        },
    )
